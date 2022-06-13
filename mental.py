
from docx import Document
import os
import natsort
class Mental:
    def __init__(self,path_ref):
        self.path_ref = path_ref
        self.ref = Document(self.path_ref)

    def getName(self):
        completedText = []
        for paragraph in self.ref.paragraphs:
            completedText.append(paragraph.text)
        refor = completedText[3]
        refor = refor.replace("…",' ').replace(".",' ')
        refor = refor.split(" ")
        delete_empty = [ele for ele in refor if ele.strip()]
        refor = ' '.join(delete_empty)
        return refor

    def getNum(self):
        completedText = []
        for paragraph in self.ref.paragraphs:
            completedText.append(paragraph.text)
        refor = completedText[4]
        refor = refor.replace("…",' ').replace(".",' ')
        refor = refor.split(" ")
        delete_empty = [ele for ele in refor if ele.strip()]
        refor = ' '.join(delete_empty)
        return refor

    def getAct(self):
        text = []
        
        for i in range(1,11):
            text.append(self.ref.tables[0].cell(i,1).paragraphs[0].text)
        return text

    def getPlace(self):
        place = []
        for i in range(1,11):
            place.append(self.ref.tables[0].cell(i,2).paragraphs[0].text)
        return place

    def getDate(self):
        date = []
        for i in range(1,11):
            date.append(self.ref.tables[0].cell(i,0).paragraphs[0].text)
        return date

    def getTime(self):
        time = []
        for i in range(1,11):
            time.append(self.ref.tables[0].cell(i,3).paragraphs[0].text)
        return time

    def getRes(self):
        res = []
        for i in range(1,11):
            res.append(self.ref.tables[0].cell(i,4).paragraphs[0].text)
        return res

    def getTotalTime(self):
        total = self.ref.tables[0].cell(11,2).text
        total = total.replace("…",' ').replace(".",' ')
        total = total.split(" ")
        delete_empty = [ele for ele in total if ele.strip()]
        total = ' '.join(delete_empty) 
        return total




if __name__ == "__main__":
    M = Mental('')
    M.getTotalTime()
